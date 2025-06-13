import fitz
from docx import Document
from docx.shared import Inches
import io
import os
import re
from datetime import datetime
from PIL import Image  # 新增PIL库用于图片格式转换


def extract_images_from_pdf(pdf_path):
    """从PDF文件中提取所有图片，并进行格式转换和验证"""
    print(f"开始处理PDF文件: {pdf_path}")
    images = []
    unsupported_images = []

    try:
        with fitz.open(pdf_path) as pdf:
            for page_num in range(len(pdf)):
                page = pdf[page_num]
                image_list = page.get_images(full=True)

                for image_index, img in enumerate(image_list, start=1):
                    xref = img[0]
                    base_image = pdf.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]

                    # 尝试使用PIL打开图片，验证图片是否有效
                    try:
                        img = Image.open(io.BytesIO(image_bytes))

                        # 转换为Word支持的格式（JPEG/PNG）
                        if image_ext.lower() not in ['jpg', 'jpeg', 'png']:
                            print(f"第 {page_num + 1} 页的第 {image_index} 张图片格式 {image_ext} 不受支持，正在转换...")
                            img_byte_arr = io.BytesIO()

                            # 如果是RGBA模式，转换为RGB以保存为JPEG
                            if img.mode == 'RGBA':
                                img = img.convert('RGB')

                            img.save(img_byte_arr, format='JPEG')
                            image_bytes = img_byte_arr.getvalue()
                            image_ext = 'jpg'

                        images.append({
                            'page': page_num + 1,
                            'index': image_index,
                            'bytes': image_bytes,
                            'ext': image_ext,
                            'valid': True
                        })
                        print(f"已提取并验证第 {page_num + 1} 页的第 {image_index} 张图片")

                    except Exception as e:
                        print(f"第 {page_num + 1} 页的第 {image_index} 张图片无法识别: {str(e)}")
                        unsupported_images.append({
                            'page': page_num + 1,
                            'index': image_index,
                            'error': str(e)
                        })

        print(f"PDF文件处理完成，共提取出 {len(images)} 张有效图片，{len(unsupported_images)} 张图片无法识别")
        return images, unsupported_images

    except Exception as e:
        print(f"处理PDF文件时出错: {str(e)}")
        return [], [{'error': str(e)}]


def create_word_with_images(images, unsupported_images, output_path):
    """创建包含提取图片的Word文档，并记录不支持的图片"""
    print(f"开始创建Word文档: {output_path}")
    doc = Document()

    # 添加文档标题
    doc.add_heading('PDF图片提取文档', 0)
    doc.add_paragraph(f'创建时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    doc.add_paragraph(f'总图片数量: {len(images) + len(unsupported_images)}')
    doc.add_paragraph(f'成功提取并插入的图片数量: {len(images)}')
    doc.add_paragraph(f'无法识别的图片数量: {len(unsupported_images)}')
    doc.add_page_break()

    # 添加成功提取的图片
    for i, img_data in enumerate(images, start=1):
        # 添加图片标题
        doc.add_heading(f'图片 {i} (来自第 {img_data["page"]} 页)', level=2)

        # 添加图片
        try:
            image_stream = io.BytesIO(img_data['bytes'])
            doc.add_picture(image_stream, width=Inches(6))

            # 添加图片描述
            doc.add_paragraph(f'图片 {i}: 来自原PDF文件的第 {img_data["page"]} 页，格式为 {img_data["ext"].upper()}')
            print(f"已添加图片 {i}/{len(images)} 到Word文档")

        except Exception as e:
            doc.add_paragraph(f'[错误] 无法插入此图片: {str(e)}')
            print(f"[错误] 无法插入图片 {i}/{len(images)}: {str(e)}")

        # 除了最后一张图片外，每张图片后添加分页符
        if i < len(images):
            doc.add_page_break()

    # 添加无法识别的图片报告
    if unsupported_images:
        doc.add_page_break()
        doc.add_heading('无法识别的图片报告', level=1)

        for i, img_error in enumerate(unsupported_images, start=1):
            doc.add_heading(f'无法识别的图片 {i}', level=2)
            doc.add_paragraph(f'位置: 第 {img_error.get("page", "未知")} 页，第 {img_error.get("index", "未知")} 张')
            doc.add_paragraph(f'错误信息: {img_error["error"]}')

    # 保存Word文档
    try:
        doc.save(output_path)
        print(f"Word文档已成功保存至: {output_path}")
        return True
    except Exception as e:
        print(f"保存Word文档时出错: {str(e)}")
        return False


def sanitize_filename(filename):
    """清理文件名，移除不允许的字符"""
    return re.sub(r'[\\/*?:"<>|]', '_', filename)


def main():
    # 获取用户输入的PDF文件路径
    pdf_path = input("请输入PDF文件的路径: ").strip()

    # 检查文件是否存在
    if not os.path.exists(pdf_path):
        print(f"错误: 文件 '{pdf_path}' 不存在")
        return

    # 检查文件是否为PDF
    if not pdf_path.lower().endswith('.pdf'):
        print(f"错误: 文件 '{pdf_path}' 不是PDF文件")
        return

    # 安装必要的依赖
    try:
        import PIL
    except ImportError:
        print("正在安装必要的依赖库...")
        os.system("pip install pillow")

    # 生成输出Word文件路径
    base_name = os.path.basename(pdf_path)
    base_name = os.path.splitext(base_name)[0]
    safe_base_name = sanitize_filename(base_name)
    output_path = f"{safe_base_name}_图片提取版.docx"

    # 提取图片并创建Word文档
    images, unsupported = extract_images_from_pdf(pdf_path)
    if images:
        create_word_with_images(images, unsupported, output_path)
        print("\n操作已完成！")
    else:
        print("未从PDF中提取到任何有效图片")


if __name__ == "__main__":
    main()